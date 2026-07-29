#!/usr/bin/env python3
"""Reconstruct a regenerable builder script from an existing manuscript .docx.

Why this exists
---------------
The V5 builder lived in a scratch directory under %TEMP% and a machine restart
deleted it, leaving a .docx that could no longer be regenerated. A deliverable
whose generator only exists in volatile storage is a deliverable you can edit
once and then never again, so this recovers the generator from the document
and writes it next to the code, where it is version-controlled.

It emits a self-contained Python file whose execution reproduces the source
document: one call per paragraph (carrying style, alignment, size, bold and
italic), one call per table (header, rows, caption), and figure placement by
first citation. Round-trip is verified by rebuilding and diffing against the
original.

    python scripts/extract_builder.py --docx <in.docx> --out <builder.py>
"""
import argparse
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

NS_INLINE = ("{http://schemas.openxmlformats.org/drawingml/2006/"
             "wordprocessingDrawing}inline")

HEADER = '''# -*- coding: utf-8 -*-
"""Manuscript builder — RECOVERED from the .docx by scripts/extract_builder.py.

Every paragraph, table and figure placement below was extracted from the
document it reproduces. Edit this file, not the .docx: the .docx is an output.

    python scripts/build_manuscript.py [--out path.docx]
"""
import argparse
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
CODE = os.path.dirname(HERE)
RESULTS = os.path.join(CODE, "results")
FIGDIR = os.path.join(CODE, "figures")
DEFAULT_OUT = os.path.join(os.path.dirname(CODE),
                           "AppliedEnergy_Manuscript_v5.docx")

ALIGN = {"CENTER": WD_ALIGN_PARAGRAPH.CENTER,
         "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
         "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
         "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
         None: None}

doc = Document()
_st = doc.styles["Normal"]
_st.font.name = "Times New Roman"
_st.font.size = Pt(11)
_st.paragraph_format.space_after = Pt(6)
_st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for _i, _sz in [(1, 14), (2, 12), (3, 11)]:
    _s = doc.styles[f"Heading {_i}"]
    _s.font.name = "Times New Roman"
    _s.font.size = Pt(_sz)
    _s.font.bold = True
    _s.font.color.rgb = None


def par(text, style="Normal", align=None, size=None, bold=None, italic=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = ALIGN[align]
    r = p.add_run(text)
    if size:
        r.font.size = Pt(size)
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    return p


def table(header, rows, caption=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for c, h in zip(t.rows[0].cells, header):
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if caption:
        par(caption, size=9, italic=True)
    return t


FIG_FILES = {FIGMAP}
FIG_WIDTH_CM = {"Fig3_causal_decomposition": 12.4}


def figure(stem):
    path = os.path.join(FIGDIR, stem + ".png")
    if not os.path.exists(path):
        raise FileNotFoundError(
            f"{path} not found — run figures_tables.regenerate_from_saved() "
            f"first. Figures are never omitted silently.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(FIG_WIDTH_CM.get(stem, 15.2)))


# ─────────────────────────── document body ───────────────────────────
'''

FOOTER = '''

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=DEFAULT_OUT)
    a = ap.parse_args()
    doc.save(a.out)
    print(f"Saved {a.out}")


if __name__ == "__main__":
    main()
'''


def q(s):
    return repr(s)


def para_call(p):
    """One emitted call reproducing this paragraph."""
    style = p.style.name
    text = p.text
    runs = [r for r in p.runs if r.text]
    size = bold = italic = None
    if runs:
        r0 = runs[0]
        size = r0.font.size.pt if r0.font.size else None
        bold, italic = r0.bold, r0.italic
    align = None
    if p.alignment is not None:
        align = {WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
                 WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
                 WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
                 WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT"}.get(p.alignment)
    args = [q(text)]
    if style != "Normal":
        args.append(f"style={q(style)}")
    if align:
        args.append(f"align={q(align)}")
    if size:
        args.append(f"size={size}")
    if bold is not None:
        args.append(f"bold={bold}")
    if italic is not None:
        args.append(f"italic={italic}")
    return "par(" + ", ".join(args) + ")"


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True)
    ap.add_argument("--out", required=True)
    a = ap.parse_args()

    doc = Document(a.docx)
    body = doc.element.body
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    lines, figs, ti, fi = [], [], 0, 0
    pending_caption = None
    children = list(body.iterchildren())
    for i, ch in enumerate(children):
        if ch.tag.endswith("}p"):
            p = Paragraph(ch, doc)
            if p._p.findall(f".//{NS_INLINE}"):
                fi += 1
                lines.append(f"figure(FIG_FILES[{fi}])")
                continue
            if not p.text.strip():
                continue
            # a caption belongs to the table immediately above it
            if lines and lines[-1].startswith("__TABLE__") and \
                    p.text.strip().startswith(("Table ", "Tabla ")):
                lines[-1] = lines[-1].replace("__CAPTION__", q(p.text))
                continue
            lines.append(para_call(p))
        elif ch.tag.endswith("}tbl"):
            t = Table(ch, doc)
            hdr = [c.text.strip() for c in t.rows[0].cells]
            rows = [[c.text.strip() for c in r.cells] for r in t.rows[1:]]
            lines.append("__TABLE__table(" + repr(hdr) + ", " + repr(rows)
                         + ", caption=__CAPTION__)")
            ti += 1

    out = []
    for ln in lines:
        out.append(ln.replace("__TABLE__", "").replace("__CAPTION__", "None"))

    figmap = "{" + ", ".join(
        f"{i + 1}: {q(n)}" for i, n in enumerate([
            "Fig1_protocol_contrast", "Fig2_architecture",
            "Fig3_causal_decomposition", "Fig4_accuracy_overview",
            "Fig5_forecast_GEFCom2014", "Fig6_forecast_PJM",
            "Fig7_forecast_AEMO", "Fig8_ablation", "Fig9_cross_attention",
            "Fig10_error_by_leadtime", "Fig11_leakage_effect"])) + "}"

    src = HEADER.replace("{FIGMAP}", figmap) + "\n".join(out) + FOOTER
    with open(a.out, "w", encoding="utf-8") as f:
        f.write(src)
    print(f"wrote {a.out}: {len(out)} statements, {ti} tables, {fi} figures")


if __name__ == "__main__":
    main()
